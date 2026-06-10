import os
import re

import docx
from docx.shared import Pt


def sanitize_filename(title):
    cleaned = re.sub(r'[\\/:*?"<>|]+', "", title)
    return cleaned.replace(" ", "_").strip("_") or "quiz"


def get_output_dir():
    return "/tmp" if os.getenv("VERCEL") else os.getcwd()


def build_question_text(question):
    return question["pertanyaan"]


def set_paragraph_spacing(paragraph, after=0, before=0):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(after)
    paragraph_format.space_before = Pt(before)


def add_indented_paragraph(document, text, left_indent_pt=18):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.left_indent = Pt(left_indent_pt)
    set_paragraph_spacing(paragraph, after=0)
    return paragraph


def split_questions_by_type(questions_list):
    pg_questions = [question for question in questions_list if question.get("tipe", "").lower() == "pg"]
    essay_questions = [question for question in questions_list if question.get("tipe", "").lower() == "esai"]
    other_questions = [
        question for question in questions_list
        if question.get("tipe", "").lower() not in {"pg", "esai"}
    ]
    return pg_questions, essay_questions, other_questions


def add_section_heading(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    set_paragraph_spacing(paragraph, after=10, before=10)
    return paragraph


def resolve_answer_option(question):
    answer_key = (question.get("kunci_jawaban") or "").strip().upper()
    if not answer_key:
        return ""

    for option in question.get("pilihan", []):
        cleaned_option = option.strip()
        prefix = cleaned_option.split(".", 1)[0].split(")", 1)[0].strip().upper()
        if prefix == answer_key or cleaned_option.upper() == answer_key:
            return cleaned_option
    return answer_key


def generate_questions_docx(title, questions_list, output_dir):
    document = docx.Document()
    document.add_heading(title, 0)
    pg_questions, essay_questions, other_questions = split_questions_by_type(questions_list)
    ordered_sections = [
        ("I. PILIHAN GANDA", pg_questions),
        ("II. ESSAY", essay_questions),
    ]
    if other_questions:
        ordered_sections.append(("III. LAINNYA", other_questions))

    question_number = 1
    for section_title, section_questions in ordered_sections:
        if not section_questions:
            continue
        add_section_heading(document, section_title)

        for question in section_questions:
            question_paragraph = document.add_paragraph()
            question_paragraph.add_run(f"{question_number}. {build_question_text(question)}")
            set_paragraph_spacing(question_paragraph, after=4)

            if question["tipe"].lower() == "pg":
                for option in question["pilihan"]:
                    add_indented_paragraph(document, option, left_indent_pt=24)
            else:
                add_indented_paragraph(
                    document,
                    "Jawaban: ...................................................................................................."
                )
            document.add_paragraph("")
            question_number += 1

    questions_filename = f"{sanitize_filename(title)}_soal.docx"
    questions_path = os.path.join(output_dir, questions_filename)
    document.save(questions_path)
    return os.path.abspath(questions_path)


def generate_answer_key_docx(title, questions_list, output_dir):
    document = docx.Document()
    document.add_heading(f"Kunci Jawaban - {title}", 0)
    has_answer_key = any((question.get("kunci_jawaban") or "").strip() for question in questions_list)
    pg_questions, essay_questions, other_questions = split_questions_by_type(questions_list)

    if not has_answer_key:
        document.add_paragraph("Dokumen ini tidak memiliki kunci jawaban karena konten dibuat sebagai survey/non-quiz.")

    ordered_sections = [
        ("I. PILIHAN GANDA", pg_questions),
        ("II. ESSAY", essay_questions),
    ]
    if other_questions:
        ordered_sections.append(("III. LAINNYA", other_questions))

    answer_number = 1
    for section_title, section_questions in ordered_sections:
        if not section_questions:
            continue
        add_section_heading(document, section_title)

        for question in section_questions:
            answer_key = question.get("kunci_jawaban", "").strip()
            explanation = (question.get("penjelasan") or "").strip()

            if question["tipe"].lower() == "pg":
                resolved_answer = resolve_answer_option(question)
                if resolved_answer:
                    answer_label = resolved_answer
                else:
                    answer_label = answer_key or "-"
            else:
                if has_answer_key:
                    answer_label = "Penilaian manual oleh guru."
                else:
                    answer_label = "Tidak digunakan."

            answer_paragraph = document.add_paragraph()
            answer_paragraph.add_run(f"{answer_number}. {answer_label}")
            set_paragraph_spacing(answer_paragraph, after=2)

            if explanation:
                add_indented_paragraph(document, f"Penjelasan: {explanation}")

            document.add_paragraph("")
            answer_number += 1

    answer_key_filename = f"{sanitize_filename(title)}_kunci_jawaban.docx"
    answer_key_path = os.path.join(output_dir, answer_key_filename)
    document.save(answer_key_path)
    return os.path.abspath(answer_key_path)


def generate_docx_file(title, questions_list):
    """Membuat dua file Word: naskah soal dan kunci jawaban."""
    if not isinstance(questions_list, list):
        raise ValueError("Data soal untuk file Word tidak valid.")

    output_dir = get_output_dir()
    questions_path = generate_questions_docx(title, questions_list, output_dir)
    answer_key_path = generate_answer_key_docx(title, questions_list, output_dir)

    return {
        "questions_file_path": questions_path,
        "answer_key_file_path": answer_key_path
    }
